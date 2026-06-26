import unittest
import io
from pathlib import Path

import fitz  # PyMuPDF
import docx

from src.services.file_processing_service import FileProcessingService, ProcessedFileResult
from src.utils.exceptions import UnsupportedFileTypeError, ExtractionError


class TestFileProcessingService(unittest.TestCase):
    def setUp(self):
        self.service = FileProcessingService()

    def test_txt_processing(self):
        txt_content = "   Hello World!  \n\n\n  This is a test.   \n"
        txt_bytes = txt_content.encode("utf-8")
        
        result = self.service.process_file(txt_bytes, file_name="test.txt")
        
        self.assertEqual(result.file_name, "test.txt")
        self.assertEqual(result.file_type, "txt")
        self.assertEqual(result.pages, 1)
        self.assertEqual(result.raw_text, txt_content)
        self.assertEqual(result.clean_text, "Hello World!\n\nThis is a test.")
        self.assertEqual(result.metadata["extension"], ".txt")
        self.assertGreater(result.metadata["file_size_kb"], 0)
        self.assertEqual(result.metadata["extraction_method"], "utf-8-reader")

    def test_pdf_processing(self):
        # Create a simple PDF dynamically in memory
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Line 1 inside PDF\n\nLine 2 inside PDF")
        pdf_bytes = doc.write()
        doc.close()

        result = self.service.process_file(pdf_bytes, file_name="test.pdf")

        self.assertEqual(result.file_name, "test.pdf")
        self.assertEqual(result.file_type, "pdf")
        self.assertEqual(result.pages, 1)
        self.assertIn("Line 1 inside PDF", result.raw_text)
        self.assertEqual(result.clean_text, "Line 1 inside PDF\nLine 2 inside PDF")
        self.assertEqual(result.metadata["extension"], ".pdf")
        self.assertEqual(result.metadata["extraction_method"], "pymupdf")

    def test_docx_processing(self):
        # Create a simple docx dynamically in memory
        doc = docx.Document()
        doc.add_paragraph("Paragraph 1 of word doc.")
        doc.add_paragraph("")
        doc.add_paragraph("Paragraph 2 of word doc.")
        
        stream = io.BytesIO()
        doc.save(stream)
        docx_bytes = stream.getvalue()

        result = self.service.process_file(docx_bytes, file_name="test.docx")

        self.assertEqual(result.file_name, "test.docx")
        self.assertEqual(result.file_type, "docx")
        self.assertEqual(result.pages, 1)
        self.assertIn("Paragraph 1 of word doc.", result.raw_text)
        self.assertEqual(result.clean_text, "Paragraph 1 of word doc.\n\nParagraph 2 of word doc.")
        self.assertEqual(result.metadata["extension"], ".docx")
        self.assertEqual(result.metadata["extraction_method"], "python-docx")

    def test_image_processing_placeholder(self):
        image_bytes = b"fake_image_bytes"
        with self.assertRaises(NotImplementedError) as context:
            self.service.process_file(image_bytes, file_name="test.png")
        
        self.assertIn("Gemini Vision in Milestone 4", str(context.exception))

    def test_unsupported_file_type(self):
        dummy_bytes = b"some content"
        with self.assertRaises(UnsupportedFileTypeError):
            self.service.process_file(dummy_bytes, file_name="test.csv")

    def test_file_like_object_processing(self):
        txt_content = "File-like object content."
        stream = io.BytesIO(txt_content.encode("utf-8"))
        
        result = self.service.process_file(stream, file_name="stream_test.txt")
        self.assertEqual(result.file_name, "stream_test.txt")
        self.assertEqual(result.file_type, "txt")
        self.assertEqual(result.clean_text, "File-like object content.")


if __name__ == "__main__":
    unittest.main()
